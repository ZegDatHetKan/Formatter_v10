#!/usr/bin/env python3
"""Deterministic `letters` formatter (first slice).

Structure reverse-engineered ONLY from the historical letter OUTPUTS and the
empty letterhead template (`assets/Template_Vuoto.docx`). No external rule files.

Pipeline:
  template (header/footer/margins) -> read input as plain text ->
  strip duplicated letterhead -> normalize text -> classify each line into a
  letter block (positional zones around the OGGETTO anchor) -> apply block style
  -> write .docx + technical report.

Uncertainty is never hidden: placeholders `[...]`, missing OGGETTO, or
unclassifiable preambles set `needs_review` in the report.
"""
from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = "Times New Roman"
BODY = 12
TITLE = 16          # central hierarchy (OGGETTO label, ritual section titles)
SUBTITLE = 14       # left-aligned subsection titles
SMALL = 10          # transmission notes / fine print
RECIPIENT_INDENT_CM = 8.5
LIST_INDENT_CM = 0.5

ASSETS = Path(__file__).resolve().parent.parent / "assets"
TEMPLATE = ASSETS / "Template_Vuoto.docx"

# --- text normalization -----------------------------------------------------
DASHES = {"–": "-", "—": "-", "−": "-"}
ZERO_WIDTH = dict.fromkeys(map(ord, "​‌‍﻿­"), None)


def normalize(text: str) -> str:
    text = text.translate(ZERO_WIDTH)
    for bad, good in DASHES.items():
        text = text.replace(bad, good)
    text = text.replace(" ", " ")
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


# --- detection patterns -----------------------------------------------------
LETTERHEAD = re.compile(
    r"^(bergamo legal|societ. tra avvocati|studio legale|www\.bergamo\.legal"
    r"|r\.e\.a\.|iscritta all.ordine)\b|\[indirizzo studio\]|\[pec studio\]",
    re.I,
)
DELIVERY = re.compile(
    r"^(trasmessa|spedita|inviata|a mezzo|raccomandata|via pec|anticipata"
    r"|e, per anticipazione)\b",
    re.I,
)
DATE_PLACE = re.compile(r"^[A-ZÀ-Ù][\wÀ-ù'\.]+,\s*(l[iì]\s*)?(\[|\d)", re.U)
OGGETTO = re.compile(r"^(oggetto)\s*:\s*(.*)$", re.I | re.S)
GREETING = re.compile(
    r"^(gentile|egregi[oa]|egr\.|spett\.?le|preg\.?m[oa]|chiar\.?m[oa]|on\.?le"
    r"|all.attenzione|signor|signora)\b.*,\s*$",
    re.I,
)
HONORIFIC = re.compile(
    r"^(spett\.?le|preg\.?m[oa]|egr\.?|gent\.?|sig\.?|all.|on\.?le|c/o|alla |al )",
    re.I,
)
# studio partner line (used by letterhead strip, never in the signature zone)
PARTNER = re.compile(r"^(avv|dott|ing|prof)\.", re.I)
# a recipient line that is *only* a courtesy prefix and must join the next line
HONORIFIC_ONLY = re.compile(
    r"^(spett\.?le|egr\.?(m[oa])?|preg\.?m[oa]|gent\.?(m[oa])?|chiar\.?m[oa]"
    r"|ill\.?m[oa]|on\.?le)\.?$",
    re.I,
)
ADDRESS = re.compile(r"^(via|viale|corso|piazza|p\.zza|str\.|loc\.)\b|\b\d{5}\b", re.I)
CONTACT = re.compile(r"^(pec|email|e-mail|fax|tel)\b", re.I)
NUMBERED = re.compile(r"^(\d+(?:\.\d+)*)([.)])\s+(.*)$", re.S)
ENUM = re.compile(r"^(\((?:[ivxlcdm]+|[a-z])\))\s+(.*)$", re.I | re.S)
BULLET = re.compile(r"^([•▪◦\-\*])\s+(.*)$")
CLOSING = re.compile(
    r"^(cordiali saluti|distinti saluti|con osservanza|con\s+i\s+migliori\s+saluti"
    r"|in attesa\b.*riscontro|resto a disposizione|resta(?:ndo)? a disposizione"
    r"|nel rimettere\b)",
    re.I,
)
SIGNATURE = re.compile(r"^(avv\.|dott\.|ing\.|prof\.|studio|bergamo legal|ammin)", re.I)
SIG_LINE = re.compile(r"^_{3,}$")
ATTACH = re.compile(r"^allegat[oi]\s*:?\s*$", re.I)
ATTACH_ITEM = re.compile(r"^all\.\s*\d", re.I)
PLACEHOLDER = re.compile(r"\[[^\]]*\]")
# ritual section titles rendered centered (16pt) when present
RITUAL = {
    "in fatto", "in diritto", "premesso che", "premesse", "tanto premesso",
    "conclusioni", "motivi", "diffida", "invita", "invita la s.v.", "chiede",
    "chiedono", "p.q.m.", "per questi motivi", "voglia",
}
# short left-aligned subsection titles (14pt)
LEFT_HEADINGS = {"fatto", "diritto", "in fatto e in diritto", "premessa"}


@dataclass
class Block:
    kind: str
    text: str
    size: int = BODY
    align: int = AL.JUSTIFY
    bold: bool = False
    italic: bool = False
    left_indent: float | None = None
    space_before: float = 0
    space_after: float = 6
    keep_with_next: bool = False
    keep_together: bool = False
    bold_label: str | None = None   # leading token rendered bold (number / label)
    label_size: int | None = None   # size of bold_label run (None = same as size)


@dataclass
class Report:
    input_path: str = ""
    output_path: str = ""
    rules_applied: list[str] = field(default_factory=list)
    blocks: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    needs_review: bool = False


# --- run/paragraph styling --------------------------------------------------
def style_run(run, size, bold=None, italic=None):
    f = run.font
    f.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)
    f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic


def emit(doc, b: Block):
    p = doc.add_paragraph()
    p.alignment = b.align
    pf = p.paragraph_format
    if b.left_indent is not None:
        pf.left_indent = Cm(b.left_indent)
    pf.space_before = Pt(b.space_before)
    pf.space_after = Pt(b.space_after)
    pf.keep_with_next = b.keep_with_next
    pf.keep_together = b.keep_together

    # Split into runs so that a leading label and any [placeholder] are bold.
    text = b.text
    if b.bold_label and text.startswith(b.bold_label):
        r = p.add_run(b.bold_label)
        style_run(r, b.label_size or b.size, bold=True, italic=b.italic)
        text = text[len(b.bold_label):]

    pos = 0
    for m in PLACEHOLDER.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            style_run(r, b.size, bold=b.bold, italic=b.italic)
        r = p.add_run(m.group(0))
        style_run(r, b.size, bold=True, italic=b.italic)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        style_run(r, b.size, bold=b.bold, italic=b.italic)
    if not p.runs:                       # empty-after-label guard
        style_run(p.add_run(""), b.size)
    return p


# --- classification ---------------------------------------------------------
def strip_letterhead(lines: list[str], rep: Report) -> list[str]:
    # Remove the duplicated studio letterhead at the very top. Partner lines
    # (`Avv. ...`) are stripped only when they trail an already-matched header
    # block (so the same names in the bottom signature are never touched, and a
    # letter addressed to a lawyer is not mistaken for a header).
    i = 0
    removed = 0
    seen_identity = False
    while i < len(lines):
        ln = lines[i]
        if ln == "":
            i += 1
            removed += 1
            continue
        if LETTERHEAD.search(ln):
            seen_identity = True
            i += 1
            removed += 1
            continue
        if seen_identity and PARTNER.match(ln):
            i += 1
            removed += 1
            continue
        break
    if removed:
        rep.rules_applied.append(f"strip_letterhead ({removed} righe)")
    return lines[i:]


def classify(lines: list[str], rep: Report) -> list[Block]:
    blocks: list[Block] = []
    # anchor: first OGGETTO line splits preamble / body
    ogg_idx = next((i for i, l in enumerate(lines) if OGGETTO.match(l)), None)
    if ogg_idx is None:
        rep.warnings.append("OGGETTO non trovato: preambolo dedotto euristicamente")
        rep.needs_review = True
        ogg_idx = next((i for i, l in enumerate(lines) if GREETING.match(l)), 0)
        body_start = ogg_idx
        preamble = lines[:ogg_idx]
        ogg_line = None
    else:
        preamble = lines[:ogg_idx]
        ogg_line = lines[ogg_idx]
        body_start = ogg_idx + 1

    # ---- preamble: delivery / date / recipient -----------------------------
    # I3: join a bare courtesy prefix ("Egr.") with the following name line.
    preamble = merge_isolated_honorifics(preamble)
    recipient_open = False
    for ln in preamble:
        if not ln:
            continue
        # I4: delivery/date are only recognized BEFORE the recipient block; a
        # delivery-like line after it is a recipient continuation, not a new line.
        if DELIVERY.match(ln) and not recipient_open:
            blocks.append(Block("delivery_method", ln, align=AL.RIGHT,
                                italic=True, space_after=2))
        elif DATE_PLACE.match(ln) and len(ln) < 40 and not recipient_open:
            blocks.append(Block("date_place", ln, align=AL.RIGHT,
                                space_before=6, space_after=10))
        else:
            bold = bool(HONORIFIC.match(ln)) or (
                ln[:1].isupper() and not (ADDRESS.search(ln) or CONTACT.match(ln)))
            blocks.append(Block("recipient", ln, align=AL.LEFT, bold=bold,
                                left_indent=RECIPIENT_INDENT_CM, space_after=2))
            recipient_open = True
    if recipient_open:
        blocks[-1].space_after = 8

    # ---- oggetto -----------------------------------------------------------
    if ogg_line is not None:
        m = OGGETTO.match(ogg_line)
        label = ogg_line[: m.start(2)].rstrip()        # "Oggetto:" / "OGGETTO:"
        blocks.append(Block("subject_line", ogg_line, size=BODY, align=AL.JUSTIFY,
                            bold=True, space_before=8, space_after=8,
                            keep_with_next=True, bold_label=label, label_size=TITLE))
        rep.rules_applied.append(
            "oggetto: etichetta 16pt bold + contenuto 12pt bold, JUSTIFY, keep_with_next")

    # ---- body --------------------------------------------------------------
    in_signature = False
    prev_lead_colon = False
    body = [l for l in lines[body_start:]]
    for j, ln in enumerate(body):
        if not ln:
            prev_lead_colon = False
            continue
        low = ln.lower().rstrip(" .:")
        first_body = (j == next((k for k, x in enumerate(body) if x), -1))

        if SIG_LINE.match(ln):
            blocks.append(Block("signature_line", ln, align=AL.RIGHT, space_after=6))
        elif ATTACH.match(ln):
            blocks.append(Block("attachments", ln, align=AL.LEFT, bold=True,
                                space_before=16, space_after=4))
            in_signature = False
        elif ATTACH_ITEM.match(ln):
            blocks.append(Block("attachment_item", ln, align=AL.LEFT,
                                left_indent=LIST_INDENT_CM, space_after=3))
        elif low in RITUAL or _is_allcaps_title(ln):
            blocks.append(Block("section_center", ln, size=TITLE, align=AL.CENTER,
                                bold=True, space_before=14, space_after=6,
                                keep_with_next=True, keep_together=True))
        elif low in LEFT_HEADINGS:
            blocks.append(Block("section_left", ln, size=SUBTITLE, align=AL.LEFT,
                                bold=True, space_before=12, space_after=6,
                                keep_with_next=True))
        elif first_body and GREETING.match(ln):
            blocks.append(Block("opening", ln, align=AL.JUSTIFY,
                                space_before=8, space_after=6))
        elif CLOSING.match(ln) or in_signature:
            italic = bool(re.match(r"^(in attesa|resta|resto|nel rimettere)", ln, re.I))
            bold = bool(SIGNATURE.match(ln)) and not italic
            blocks.append(Block("closing_signature", ln, align=AL.RIGHT,
                                bold=bold, italic=italic,
                                space_before=12 if not in_signature else 4,
                                space_after=4))
            in_signature = True
        elif (m := NUMBERED.match(ln)):
            blocks.append(Block("numbered", ln, align=AL.JUSTIFY, space_after=6,
                                bold_label=m.group(1) + m.group(2) + " "))
        elif (m := ENUM.match(ln)):
            blocks.append(Block("enum_item", ln, align=AL.JUSTIFY,
                                left_indent=LIST_INDENT_CM, space_after=4))
        elif BULLET.match(ln):
            blocks.append(Block("bullet_item", ln, align=AL.JUSTIFY,
                                left_indent=LIST_INDENT_CM, space_after=4))
        elif prev_lead_colon and ln.endswith(";"):
            blocks.append(Block("list_item", ln, align=AL.JUSTIFY,
                                left_indent=LIST_INDENT_CM, space_after=4))
        else:
            blocks.append(Block("body", ln, align=AL.JUSTIFY, space_after=6))

        prev_lead_colon = ln.endswith(":")

    for b in blocks:
        rep.blocks.append(f"{b.kind}: {b.text[:55]!r}")
        if PLACEHOLDER.search(b.text):
            rep.needs_review = True
    if any(PLACEHOLDER.search(b.text) for b in blocks):
        rep.warnings.append("placeholder [...] presenti: compilazione manuale richiesta")
    return blocks


def merge_isolated_honorifics(lines: list[str]) -> list[str]:
    """Join a line that is only a courtesy prefix with the next line (I3).

    `Egr.` / `Spett.le` alone -> `Egr. Sig.ra ...` / `Spett.le <name>`.
    """
    out: list[str] = []
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln and HONORIFIC_ONLY.match(ln) and i + 1 < len(lines) and lines[i + 1]:
            out.append(f"{ln.rstrip()} {lines[i + 1].lstrip()}")
            i += 2
        else:
            out.append(ln)
            i += 1
    return out


def _is_allcaps_title(ln: str) -> bool:
    letters = [c for c in ln if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters) and len(ln) < 60


# --- entry point ------------------------------------------------------------
def format_letter(input_path: str, output_path: str,
                  template_path: str | None = None) -> Report:
    rep = Report(input_path=input_path, output_path=output_path)
    tpl = Path(template_path) if template_path else TEMPLATE
    src = Document(input_path)
    # Inputs vary: some use real paragraph breaks, others pack the whole letter
    # into one paragraph separated by newline characters. Split on both.
    lines = [normalize(sub) for p in src.paragraphs for sub in p.text.split("\n")]
    rep.rules_applied.append("normalize: dash->'-', nbsp/zero-width/collapse-spaces")

    lines = strip_letterhead(lines, rep)
    blocks = classify(lines, rep)

    out = Document(str(tpl))
    # remove the template's placeholder empty body paragraph(s)
    for p in list(out.paragraphs):
        p._element.getparent().remove(p._element)
    for b in blocks:
        emit(out, b)
    out.save(output_path)
    rep.rules_applied.append(f"template: {tpl.name} (intestazione preservata)")
    return rep


def render_report(rep: Report) -> str:
    lines = [
        "# Letters Formatter Report",
        "",
        "- mode: letters",
        f"- input: {rep.input_path}",
        f"- output: {rep.output_path}",
        f"- needs_review: {str(rep.needs_review).lower()}",
        "",
        "## Rules applied",
        *[f"- {r}" for r in rep.rules_applied],
        "",
        f"## Blocks detected ({len(rep.blocks)})",
        *[f"- {b}" for b in rep.blocks],
        "",
        "## Warnings",
        *([f"- {w}" for w in rep.warnings] or ["- none"]),
        "",
    ]
    return "\n".join(lines)


if __name__ == "__main__":
    inp, outp = sys.argv[1], sys.argv[2]
    r = format_letter(inp, outp)
    rep_path = Path(outp).with_suffix("").as_posix() + "_report.md"
    Path(rep_path).write_text(render_report(r), encoding="utf-8")
    print(f"OK {outp}  needs_review={r.needs_review}  blocks={len(r.blocks)}")
