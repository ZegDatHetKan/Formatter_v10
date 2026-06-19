#!/usr/bin/env python3
"""Test suite for the deterministic letters formatter.

Run with the project venv:
    WebApp/.venv/bin/python -m pytest tests/test_letters.py -q

Two layers:
  - CONTRACT tests: invariants the formatter must always satisfy.
  - KNOWN-ISSUE tests (xfail): encode feedback-derived bugs I1/I2/I3 so that a
    future fix turns them into XPASS and flags the regression boundary.
"""
from __future__ import annotations

import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.shared import Emu

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
from formatters import letters  # noqa: E402


# --- helpers ----------------------------------------------------------------
def build_input(tmp_path: Path, lines, single_blob=False) -> Path:
    d = Document()
    if single_blob:
        d.add_paragraph("\n".join(lines))
    else:
        for ln in lines:
            d.add_paragraph(ln)
    p = tmp_path / "in.docx"
    d.save(p)
    return p


def fmt(tmp_path: Path, lines, single_blob=False):
    inp = build_input(tmp_path, lines, single_blob)
    outp = tmp_path / "out.docx"
    rep = letters.format_letter(str(inp), str(outp))
    return Document(str(outp)), rep


def dom_size(p):
    c = Counter(round(r.font.size.pt, 1) for r in p.runs if r.text.strip() and r.font.size)
    return c.most_common(1)[0][0] if c else None


def cm(v):
    return None if v is None else round(Emu(v).cm, 2)


PEC = [
    "Trasmessa a mezzo PEC",
    "Bergamo, lì [DA INSERIRE: data]",
    "Spett.le",
    "Gamma Servizi S.r.l.",
    "Via Europa, 41 – 24035 Curno (BG)",
    "OGGETTO: Contestazione fattura n. 88/2026",
    "Spett.le Società,",
    "lo scrivente Studio assiste la Delta S.a.s.",
    "La pretesa è integralmente contestata.",
    "In attesa di cortese riscontro, si porgono distinti saluti.",
    "Avv. Matteo Bertocchi",
]


# --- pure-function unit tests ----------------------------------------------
def test_normalize_dashes_and_spaces():
    assert letters.normalize("a – b — c") == "a - b - c"
    assert letters.normalize("x y") == "x y"
    assert letters.normalize("a   b\t c") == "a b c"
    assert letters.normalize("zero​width") == "zerowidth"


def test_strip_letterhead_removes_studio_block():
    rep = letters.Report()
    lines = ["BERGAMO LEGAL", "Società tra Avvocati s.r.l.", "www.bergamo.legal",
             "Trasmessa a mezzo PEC", "Spett.le"]
    out = letters.strip_letterhead(lines, rep)
    assert out[0] == "Trasmessa a mezzo PEC"
    assert any("strip_letterhead" in r for r in rep.rules_applied)


# --- contract tests on the full pipeline -----------------------------------
def test_runs_and_is_valid_docx(tmp_path):
    doc, rep = fmt(tmp_path, PEC)
    assert len(doc.paragraphs) > 0
    assert rep.output_path.endswith("out.docx")


def test_letterhead_and_margins_preserved(tmp_path):
    doc, _ = fmt(tmp_path, PEC)
    s = doc.sections[0]
    assert (cm(s.top_margin), cm(s.bottom_margin), cm(s.left_margin), cm(s.right_margin)) \
        == (5.91, 4.54, 2.0, 2.0)
    header_txt = " ".join(p.text for p in s.header.paragraphs)
    assert "BERGAMO LEGAL" in header_txt
    media = [str(p.partname) for p in doc.part.package.iter_parts() if "media" in str(p.partname)]
    assert any("image1" in m for m in media)


def test_every_run_has_explicit_times_new_roman_and_size(tmp_path):
    doc, _ = fmt(tmp_path, PEC)
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text.strip():
                assert r.font.name == "Times New Roman"
                assert r.font.size is not None


def test_no_en_or_em_dash_in_output(tmp_path):
    doc, _ = fmt(tmp_path, PEC)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "–" not in full and "—" not in full


def test_text_is_preserved(tmp_path):
    # every input content token (minus letterhead vocabulary) survives
    doc, _ = fmt(tmp_path, PEC)
    out = " ".join(p.text for p in doc.paragraphs).replace("–", "-").lower()
    out_tokens = Counter(re.findall(r"\w+", out))
    skip = {"bergamo", "legal", "società", "avvocati", "s", "r", "l"}
    for tok in re.findall(r"\w+", " ".join(PEC).replace("–", "-").lower()):
        if tok in skip:
            continue
        assert out_tokens[tok] >= 1, f"token perso: {tok}"


def test_oggetto_is_single_paragraph(tmp_path):
    # regression guard for feedback 027/009 ("Oggetto: a capo / centrale")
    doc, _ = fmt(tmp_path, PEC)
    ogg = [p for p in doc.paragraphs if p.text.strip().lower().startswith("oggetto")]
    assert len(ogg) == 1
    # label + content live in the SAME paragraph, not split / not centered
    assert "fattura" in ogg[0].text.lower()
    assert ogg[0].alignment != AL.CENTER


def test_recipient_block_indented_right(tmp_path):
    doc, _ = fmt(tmp_path, PEC)
    recip = [p for p in doc.paragraphs if "Gamma Servizi" in p.text]
    assert recip and cm(recip[0].paragraph_format.left_indent) == 8.5
    assert recip[0].alignment == AL.LEFT


def test_date_and_signature_right_aligned(tmp_path):
    doc, _ = fmt(tmp_path, PEC)
    date = [p for p in doc.paragraphs if p.text.startswith("Bergamo, lì")][0]
    sign = [p for p in doc.paragraphs if "Avv. Matteo Bertocchi" in p.text][-1]
    assert date.alignment == AL.RIGHT
    assert sign.alignment == AL.RIGHT


def test_placeholders_trigger_needs_review(tmp_path):
    _, rep = fmt(tmp_path, PEC)               # contains [DA INSERIRE: data]
    assert rep.needs_review is True
    assert any("placeholder" in w for w in rep.warnings)


def test_clean_letter_not_flagged(tmp_path):
    clean = [
        "Oggetto: Aggiornamento pratica",
        "Gentile Stefania,",
        "la ringrazio per il riscontro.",
        "Cordiali saluti",
        "Avv. Matteo Bertocchi",
    ]
    _, rep = fmt(tmp_path, clean)
    assert rep.needs_review is False


def test_single_blob_input_is_split(tmp_path):
    # whole letter in one paragraph separated by '\n' must still be segmented
    doc, _ = fmt(tmp_path, PEC, single_blob=True)
    assert len([p for p in doc.paragraphs if p.text.strip()]) >= 8


def test_missing_oggetto_sets_needs_review(tmp_path):
    no_ogg = ["Gentile Stefania,", "un breve aggiornamento.", "Cordiali saluti",
              "Avv. Matteo Bertocchi"]
    _, rep = fmt(tmp_path, no_ogg)
    assert rep.needs_review is True
    assert any("OGGETTO" in w for w in rep.warnings)


# --- feedback-derived fixes (I1/I2/I3) — now expected to PASS -----------------
def test_oggetto_content_is_body_size(tmp_path):
    doc, _ = fmt(tmp_path, PEC)
    ogg = [p for p in doc.paragraphs if p.text.lower().startswith("oggetto")][0]
    content_sizes = {round(r.font.size.pt, 1) for r in ogg.runs
                     if r.text.strip() and not r.text.lower().startswith("oggetto")}
    assert content_sizes == {12.0}


def test_partner_lines_stripped_from_top(tmp_path):
    lines = ["BERGAMO LEGAL", "Società tra Avvocati s.r.l.", "www.bergamo.legal",
             "Avv. Matteo Bertocchi", "Avv. Daiana Chiappa",
             "Trasmessa a mezzo PEC"] + PEC[1:]
    doc, _ = fmt(tmp_path, lines)
    # the first body paragraph must NOT be a leaked partner line at indent 8.5
    first = next(p for p in doc.paragraphs if p.text.strip())
    assert not (cm(first.paragraph_format.left_indent) == 8.5
                and first.text.strip() == "Avv. Matteo Bertocchi")


def test_isolated_honorific_merged_with_name(tmp_path):
    lines = ["Bergamo, lì [DA INSERIRE: data]", "Egr.", "Sig.ra Simona Locafaro",
             "Via Recupero, 27 – 70044 Polignano a Mare (BA)",
             "OGGETTO: Diffida", "Egregia Signora,", "testo.", "Distinti saluti",
             "Avv. Matteo Bertocchi"]
    doc, _ = fmt(tmp_path, lines)
    assert any(p.text.strip() == "Egr. Sig.ra Simona Locafaro" for p in doc.paragraphs)
