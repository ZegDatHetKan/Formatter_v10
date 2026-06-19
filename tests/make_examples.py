#!/usr/bin/env python3
"""Generate 5 synthetic example letters (raw, unformatted) for the formatter.

These are NOT from the historical corpus — they are fresh inputs written to
exercise the deterministic letters formatter end to end. Saved to examples/.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

OUT = Path(__file__).resolve().parent.parent / "examples"


def make_paragraphs(name: str, lines: list[str]) -> Path:
    """Each logical line becomes a real paragraph (common input shape)."""
    d = Document()
    for ln in lines:
        d.add_paragraph(ln)
    p = OUT / name
    d.save(p)
    return p


def make_single_blob(name: str, lines: list[str]) -> Path:
    """Whole letter packed in ONE paragraph separated by newlines (the other shape)."""
    d = Document()
    d.add_paragraph("\n".join(lines))
    p = OUT / name
    d.save(p)
    return p


# 1) Diffida: letterhead to strip, sections Fatto/Diritto, (i)(ii), DIFFIDA, allegati.
DIFFIDA = [
    "BERGAMO LEGAL",
    "Società tra Avvocati s.r.l.",
    "www.bergamo.legal",
    "Avv. Matteo Bertocchi",
    "Avv. Daiana Chiappa",
    "A mezzo Raccomandata A/R, anticipata via email",
    "Bergamo, lì [DA INSERIRE: data]",
    "Spett.le",
    "Beta Costruzioni S.r.l.",
    "Via Roma, 12 – 24100 Bergamo (BG)",
    "PEC: beta.costruzioni@pec.it",
    "OGGETTO: Diffida ad adempiere – pagamento fatture insolute n. 14/2026 e n. 21/2026",
    "Egregi Signori,",
    "lo Studio scrivente agisce in nome e per conto della Alfa S.p.A. in relazione ai fatti che seguono.",
    "Fatto",
    "1. In data 3 marzo 2026 la Alfa S.p.A. emetteva regolare fattura n. 14/2026 per l'importo di € 12.500,00.",
    "2. Nonostante i solleciti, la S.V. non ha provveduto al pagamento delle seguenti fatture:",
    "(i) fattura n. 14/2026, scaduta il 2 aprile 2026;",
    "(ii) fattura n. 21/2026, scaduta il 30 aprile 2026.",
    "Diritto",
    "3. L'inadempimento integra una grave violazione degli artt. 1218 e 1453 c.c.",
    "Tutto ciò premesso e considerato, lo Studio scrivente",
    "DIFFIDA",
    "la S.V., come sopra individuata, a voler provvedere entro 15 giorni al pagamento di quanto dovuto.",
    "In difetto, si procederà nelle competenti sedi giudiziarie con aggravio di spese.",
    "In attesa di un cortese e sollecito riscontro, si porgono distinti saluti.",
    "Avv. Matteo Bertocchi",
    "Avv. Daiana Chiappa",
    "Allegati:",
    "All. 1 – copia fattura n. 14/2026;",
    "All. 2 – copia fattura n. 21/2026.",
]

# 2) PEC breve: delivery a destra, destinatario, oggetto inline, firma studio.
PEC = [
    "Trasmessa a mezzo PEC",
    "Bergamo, lì [DA INSERIRE: data]",
    "Spett.le",
    "Gamma Servizi S.r.l.",
    "Via Europa, 41 – 24035 Curno (BG)",
    "PEC: gamma.servizi@pec.it",
    "OGGETTO: Contestazione fattura n. 88/2026",
    "Spett.le Società,",
    "lo scrivente Studio assiste la Delta S.a.s. in relazione alla fattura in oggetto.",
    "La pretesa è integralmente contestata in quanto riferita a prestazioni mai rese.",
    "Si invita pertanto Codesta Società a emettere nota di credito a storno entro 10 giorni.",
    "In attesa di cortese riscontro, si porgono distinti saluti.",
    "Bergamo Legal Società tra Avvocati S.r.l.",
    "Avv. Matteo Bertocchi",
]

# 3) Comunicazione operativa: SINGLE-paragraph input, greeting, bullet list, saluti.
COMUNICAZIONE = [
    "Oggetto: Aggiornamento pratica e documenti richiesti",
    "Gentile Stefania,",
    "la ringrazio per il riscontro fornito nei giorni scorsi.",
    "Le confermo che ho provveduto a trasmettere la seguente documentazione:",
    "• visura camerale aggiornata;",
    "• certificato di iscrizione all'Albo;",
    "• copia di un documento di identità.",
    "Resto a disposizione per ogni ulteriore chiarimento.",
    "Cordiali saluti",
    "Avv. Matteo Bertocchi",
]

# 4) Riscontro: typed letterhead to strip, place/date, plain body, closing right.
RISCONTRO = [
    "Studio Legale Bergamo Legal",
    "[Indirizzo Studio] – [CAP] Bergamo – PEC: [pec studio]",
    "Bergamo, lì [DA INSERIRE: data]",
    "Egr.",
    "Sig. Marco De Luca",
    "Via Verdi, 9 – 24121 Bergamo (BG)",
    "OGGETTO: Riscontro alla Vostra del 12 maggio 2026",
    "Egregio Signore,",
    "in riscontro alla Vostra comunicazione in oggetto, si rappresenta quanto segue.",
    "La richiesta non può essere accolta nei termini formulati, per le ragioni di seguito esposte.",
    "Si resta a disposizione per un incontro chiarificatore.",
    "Distinti saluti",
    "Avv. Daiana Chiappa",
]

# 5) Recesso: raccomandata, numbered paragraphs, placeholders, signature.
RECESSO = [
    "A mezzo Raccomandata A/R",
    "Bergamo, lì [DA INSERIRE: data]",
    "Spett.le",
    "Avantgarde Immobiliare S.r.l.",
    "Via Montenapoleone, 8 – 20121 Milano (MI)",
    "OGGETTO: Recesso per gravi motivi dal contratto di sublocazione del [DA INSERIRE: data contratto]",
    "Spett.le Società,",
    "con la presente il sottoscritto comunica formalmente quanto segue.",
    "1. Il contratto di sublocazione in oggetto prevede la facoltà di recesso per gravi motivi.",
    "2. Sono sopravvenuti gravi motivi che impongono la cessazione anticipata del rapporto.",
    "Tanto premesso, il sottoscritto",
    "RECEDE",
    "dal contratto con effetto dal [DA INSERIRE: data], nel rispetto del preavviso contrattuale.",
    "Con osservanza,",
    "[DA INSERIRE: nome e cognome]",
]


def main() -> None:
    OUT.mkdir(exist_ok=True)
    made = [
        make_paragraphs("esempio_1_diffida.docx", DIFFIDA),
        make_paragraphs("esempio_2_pec.docx", PEC),
        make_single_blob("esempio_3_comunicazione.docx", COMUNICAZIONE),
        make_paragraphs("esempio_4_riscontro.docx", RISCONTRO),
        make_paragraphs("esempio_5_recesso.docx", RECESSO),
    ]
    for p in made:
        print("created", p.relative_to(OUT.parent))


if __name__ == "__main__":
    main()
