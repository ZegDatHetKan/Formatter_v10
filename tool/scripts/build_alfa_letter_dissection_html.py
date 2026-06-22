#!/usr/bin/env python3
"""Build a client-facing dissection page for the Alfa synthetic letter.

The page is generated from the actual formatted DOCX, so it shows measured Word
properties rather than hand-written guesses. The same shape can be reused for
another document family by changing the role map and input DOCX.
"""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu


TOOL_ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = TOOL_ROOT.parent
DOCX = REPO_ROOT / "out" / "examples" / "esempio_1_diffida.docx"
SOURCE_INPUT = REPO_ROOT / "examples" / "esempio_1_diffida.docx"
REPORT_SOURCE = REPO_ROOT / "out" / "examples" / "esempio_1_diffida_report.md"
OUT = TOOL_ROOT / "output" / "html" / "alfa_letter_dissection.html"
REPORT = TOOL_ROOT / "output" / "reports" / "alfa_letter_dissection_report.md"

ALIGN = {
    None: "-",
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
}


@dataclass(frozen=True)
class Role:
    kind: str
    label: str
    detection: str
    client_check: str


ROLE_BY_INDEX = {
    0: Role("delivery", "Mezzo di invio", "Prima riga del preambolo; pattern delivery_method.", "Confermare allineamento a destra e corsivo."),
    1: Role("date", "Luogo e data", "Preambolo breve con luogo + data/placeholder.", "Confermare posizione a destra e gestione placeholder."),
    2: Role("recipient", "Destinatario - nome", "Recipient block prima di OGGETTO; onorifico fuso col nome.", "Confermare rientro dal centro pagina e grassetto sul nome."),
    3: Role("recipient", "Destinatario - indirizzo", "Continuazione recipient block.", "Confermare rientro e assenza di grassetto."),
    4: Role("recipient", "Destinatario - PEC", "Continuazione recipient block / recapito.", "Confermare spazio finale del blocco destinatario."),
    5: Role("subject", "Oggetto", "Ancora strutturale OGGETTO.", "Confermare etichetta 16 pt + contenuto 12 pt nello stesso paragrafo."),
    6: Role("opening", "Formula di apertura", "Prima riga del corpo dopo OGGETTO.", "Confermare se deve essere giustificata o sinistra."),
    7: Role("body", "Corpo - premessa", "Default body paragraph.", "Confermare font, giustificazione e spaziatura."),
    8: Role("section_left", "Sezione sinistra", "Titolo breve Fatto/Diritto.", "Confermare 14 pt bold, keep_with_next."),
    9: Role("numbered", "Paragrafo numerato", "Pattern N. testo.", "Confermare prefisso numero in grassetto."),
    10: Role("numbered", "Paragrafo numerato", "Pattern N. testo.", "Confermare prefisso numero in grassetto."),
    11: Role("enum", "Sotto-elenco", "Pattern (i) / (ii).", "Confermare rientro 0.5 cm."),
    12: Role("enum", "Sotto-elenco", "Pattern (i) / (ii).", "Confermare rientro 0.5 cm."),
    13: Role("section_left", "Sezione sinistra", "Titolo breve Fatto/Diritto.", "Confermare 14 pt bold, keep_with_next."),
    14: Role("numbered", "Paragrafo numerato", "Pattern N. testo.", "Confermare prefisso numero in grassetto."),
    15: Role("body", "Formula transizione", "Default body paragraph.", "Confermare se resta body o diventa formula speciale."),
    16: Role("section_center", "Titolo rituale centrale", "Titolo rituale DIFFIDA.", "Confermare 16 pt bold, centrato, keep_together."),
    17: Role("body", "Corpo - intimazione", "Default body paragraph.", "Confermare body standard."),
    18: Role("body", "Corpo - conseguenza", "Default body paragraph.", "Confermare body standard."),
    19: Role("closing", "Chiusura", "Formula di chiusura.", "Confermare destra e corsivo."),
    20: Role("signature", "Firma", "Riga Avv. in signature block.", "Confermare destra e grassetto."),
    21: Role("signature", "Firma", "Riga Avv. in signature block.", "Confermare destra e grassetto."),
    22: Role("attachments", "Allegati - label", "Label Allegati.", "Confermare spazio prima e grassetto."),
    23: Role("attachment_item", "Allegato", "Pattern All. N.", "Confermare rientro 0.5 cm."),
    24: Role("attachment_item", "Allegato", "Pattern All. N.", "Confermare rientro 0.5 cm."),
}

COLORS = {
    "template": "#e8eef7",
    "delivery": "#d7f3ec",
    "date": "#fff0bf",
    "recipient": "#fde1d9",
    "subject": "#dfe5ff",
    "opening": "#efe3ff",
    "body": "#eef7df",
    "section_left": "#f7e6c6",
    "section_center": "#ffd9e8",
    "numbered": "#d8ecff",
    "enum": "#dff7ff",
    "closing": "#ffe5c8",
    "signature": "#f3dfc8",
    "attachments": "#e2e3e5",
    "attachment_item": "#eef0f2",
}


def cm(value):
    return None if value is None else round(Emu(value).cm, 2)


def pt(value):
    return None if value is None else round(value.pt, 1)


def fmt(value, unit: str = "") -> str:
    if value is None:
        return "-"
    return f"{value}{unit}"


def bool_label(value) -> str:
    if value is True:
        return "yes"
    if value is False:
        return "no"
    return "-"


def run_specs(paragraph):
    out = []
    for run in paragraph.runs:
        if not run.text:
            continue
        font = run.font
        out.append({
            "text": run.text,
            "font": font.name or "-",
            "size": pt(font.size),
            "bold": font.bold,
            "italic": font.italic,
            "underline": font.underline,
        })
    return out


def dominant_run(runs):
    if not runs:
        return {"font": "-", "size": None, "bold": None, "italic": None, "underline": None}
    fonts = Counter((r["font"], r["size"], r["bold"], r["italic"], r["underline"]) for r in runs)
    font, size, bold, italic, underline = fonts.most_common(1)[0][0]
    return {"font": font, "size": size, "bold": bold, "italic": italic, "underline": underline}


def para_measure(index, paragraph):
    fmtp = paragraph.paragraph_format
    runs = run_specs(paragraph)
    dom = dominant_run(runs)
    return {
        "index": index,
        "text": paragraph.text,
        "runs": runs,
        "font": dom["font"],
        "size": dom["size"],
        "bold": dom["bold"],
        "italic": dom["italic"],
        "underline": dom["underline"],
        "align": ALIGN.get(paragraph.alignment, str(paragraph.alignment)),
        "left_indent": cm(fmtp.left_indent),
        "right_indent": cm(fmtp.right_indent),
        "first_line_indent": cm(fmtp.first_line_indent),
        "space_before": pt(fmtp.space_before),
        "space_after": pt(fmtp.space_after),
        "line_spacing": fmtp.line_spacing,
        "keep_with_next": fmtp.keep_with_next,
        "keep_together": fmtp.keep_together,
    }


def section_measure(doc):
    section = doc.sections[0]
    return {
        "page_width": cm(section.page_width),
        "page_height": cm(section.page_height),
        "top": cm(section.top_margin),
        "bottom": cm(section.bottom_margin),
        "left": cm(section.left_margin),
        "right": cm(section.right_margin),
        "header": cm(section.header_distance),
        "footer": cm(section.footer_distance),
        "header_text": [p.text for p in section.header.paragraphs if p.text.strip()],
        "footer_text": [p.text for p in section.footer.paragraphs if p.text.strip()],
        "media": [str(part.partname) for part in doc.part.package.iter_parts() if "media" in str(part.partname)],
    }


def render_runs(runs) -> str:
    if not runs:
        return "<em>Nessun run testuale</em>"
    items = []
    for i, run in enumerate(runs, start=1):
        txt = escape(run["text"])
        if len(txt) > 90:
            txt = txt[:87] + "..."
        items.append(
            "<tr>"
            f"<td>{i}</td>"
            f"<td>{txt}</td>"
            f"<td>{escape(str(run['font']))}</td>"
            f"<td>{fmt(run['size'], ' pt')}</td>"
            f"<td>{bool_label(run['bold'])}</td>"
            f"<td>{bool_label(run['italic'])}</td>"
            f"<td>{bool_label(run['underline'])}</td>"
            "</tr>"
        )
    return (
        "<table class=\"run-table\"><thead><tr><th>#</th><th>Testo run</th><th>Font</th>"
        "<th>Size</th><th>Bold</th><th>Italic</th><th>Underline</th></tr></thead><tbody>"
        + "".join(items)
        + "</tbody></table>"
    )


def render_paragraph(measure) -> str:
    role = ROLE_BY_INDEX[measure["index"]]
    text = escape(measure["text"])
    color = COLORS[role.kind]
    specs = [
        ("Ruolo", role.label),
        ("Detection", role.detection),
        ("Allineamento", measure["align"]),
        ("Font dominante", measure["font"]),
        ("Dimensione dominante", fmt(measure["size"], " pt")),
        ("Bold dominante", bool_label(measure["bold"])),
        ("Italic dominante", bool_label(measure["italic"])),
        ("Rientro sx", fmt(measure["left_indent"], " cm")),
        ("Rientro dx", fmt(measure["right_indent"], " cm")),
        ("Prima riga", fmt(measure["first_line_indent"], " cm")),
        ("Spazio prima", fmt(measure["space_before"], " pt")),
        ("Spazio dopo", fmt(measure["space_after"], " pt")),
        ("Interlinea", fmt(measure["line_spacing"])),
        ("Keep with next", bool_label(measure["keep_with_next"])),
        ("Keep together", bool_label(measure["keep_together"])),
        ("Da confermare", role.client_check),
    ]
    spec_html = "".join(f"<dt>{escape(k)}</dt><dd>{escape(str(v))}</dd>" for k, v in specs)
    return f"""
      <article class="para-card" style="--role-color: {color}">
        <div class="para-number">P{measure['index']:02d}</div>
        <div class="para-body">
          <div class="para-meta"><strong>{escape(role.kind)}</strong> · {escape(role.label)}</div>
          <p class="para-text">{text}</p>
          <dl class="spec-grid">{spec_html}</dl>
          <details>
            <summary>Run interni e formattazione mista</summary>
            {render_runs(measure['runs'])}
          </details>
        </div>
      </article>
    """


def render_legend() -> str:
    rows = []
    seen = set()
    for role in ROLE_BY_INDEX.values():
        if role.kind in seen:
            continue
        seen.add(role.kind)
        rows.append(
            f"<div class=\"legend-row\"><span style=\"background:{COLORS[role.kind]}\"></span>"
            f"<strong>{escape(role.kind)}</strong><p>{escape(role.label)}. {escape(role.client_check)}</p></div>"
        )
    return "".join(rows)


def render_html() -> str:
    doc = Document(str(DOCX))
    section = section_measure(doc)
    paragraphs = [para_measure(i, p) for i, p in enumerate(doc.paragraphs)]
    generated = datetime.now(timezone.utc).isoformat(timespec="seconds")
    cards = "\n".join(render_paragraph(p) for p in paragraphs)
    header_lines = "".join(f"<li>{escape(x)}</li>" for x in section["header_text"])
    footer_lines = "".join(f"<li>{escape(x)}</li>" for x in section["footer_text"])
    media_lines = "".join(f"<li>{escape(x)}</li>" for x in section["media"]) or "<li>-</li>"
    return f"""<!doctype html>
<html lang="it">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Vivisezione lettera Alfa - Bergamo Legal</title>
  <style>
    :root {{
      --ink: #1f2933;
      --muted: #667085;
      --line: #cfd7e3;
      --paper: #ffffff;
      --bg: #f4f6f8;
    }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; background: var(--bg); color: var(--ink); font-family: Arial, Helvetica, sans-serif; line-height: 1.45; }}
    main {{ max-width: 1420px; margin: 0 auto; padding: 28px 24px 52px; }}
    h1 {{ margin: 0 0 8px; font-size: 30px; letter-spacing: 0; }}
    h2 {{ margin: 30px 0 12px; font-size: 20px; letter-spacing: 0; }}
    p {{ margin: 0 0 12px; color: var(--muted); }}
    code {{ background: #eef2f6; border: 1px solid #d8e0eb; border-radius: 4px; padding: 1px 4px; }}
    .top-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); gap: 14px; }}
    .panel {{ background: var(--paper); border: 1px solid var(--line); border-radius: 8px; padding: 14px; }}
    .panel h3 {{ margin: 0 0 8px; font-size: 15px; }}
    .panel ul {{ margin: 8px 0 0; padding-left: 18px; }}
    .layout {{ display: grid; grid-template-columns: minmax(360px, 1fr) minmax(360px, 470px); gap: 18px; align-items: start; }}
    .letter-stack {{ display: grid; gap: 9px; }}
    .para-card {{ display: grid; grid-template-columns: 52px 1fr; background: var(--paper); border: 1px solid var(--line); border-left: 12px solid var(--role-color); border-radius: 8px; overflow: hidden; }}
    .para-number {{ background: rgba(31,41,51,.05); padding: 10px 8px; font-weight: 700; font-size: 13px; text-align: center; }}
    .para-body {{ padding: 10px 12px; }}
    .para-meta {{ margin-bottom: 5px; color: #344054; font-size: 13px; }}
    .para-text {{ margin: 0 0 10px; color: var(--ink); font-family: "Times New Roman", Times, serif; font-size: 16px; }}
    .spec-grid {{ display: grid; grid-template-columns: minmax(130px, 180px) 1fr; gap: 0; margin: 0; border: 1px solid #e5ebf2; border-bottom: 0; }}
    .spec-grid dt, .spec-grid dd {{ margin: 0; padding: 6px 8px; border-bottom: 1px solid #e5ebf2; font-size: 13px; }}
    .spec-grid dt {{ background: #f7f9fb; font-weight: 700; color: #475467; }}
    details {{ margin-top: 9px; }}
    summary {{ cursor: pointer; color: #344054; font-size: 13px; font-weight: 700; }}
    .run-table {{ width: 100%; border-collapse: collapse; margin-top: 8px; font-size: 12px; }}
    .run-table th, .run-table td {{ border: 1px solid #e5ebf2; padding: 5px 6px; vertical-align: top; }}
    .run-table th {{ background: #eef2f6; }}
    .legend {{ position: sticky; top: 14px; }}
    .legend-row {{ display: grid; grid-template-columns: 22px 1fr; gap: 9px; padding: 8px 0; border-bottom: 1px solid #edf1f5; }}
    .legend-row span {{ width: 22px; height: 22px; border: 1px solid rgba(31,41,51,.2); border-radius: 4px; }}
    .legend-row strong {{ display: block; font-size: 13px; }}
    .legend-row p {{ margin: 2px 0 0; font-size: 13px; }}
    .review-list li {{ margin-bottom: 8px; }}
    @media (max-width: 980px) {{ main {{ padding: 18px 12px 36px; }} .layout {{ grid-template-columns: 1fr; }} .legend {{ position: static; }} }}
  </style>
</head>
<body>
<main>
  <h1>Vivisezione della lettera Alfa</h1>
  <p>
    Pagina di verifica cliente per la lettera fittizia <strong>Alfa S.p.A. / Beta Costruzioni S.r.l.</strong>,
    generata dal formatter <code>letters</code> sopra la pagina intestata <code>Template_Vuoto.docx</code>.
    Ogni paragrafo mostra ruolo, colore, detection rule e misure Word effettive.
  </p>

  <section class="top-grid">
    <div class="panel">
      <h3>Documento analizzato</h3>
      <ul>
        <li>Input grezzo: <code>{escape(str(SOURCE_INPUT.relative_to(REPO_ROOT)))}</code></li>
        <li>DOCX formattato: <code>{escape(str(DOCX.relative_to(REPO_ROOT)))}</code></li>
        <li>Report formatter: <code>{escape(str(REPORT_SOURCE.relative_to(REPO_ROOT)))}</code></li>
      </ul>
    </div>
    <div class="panel">
      <h3>Pagina intestata e margini</h3>
      <ul>
        <li>Formato pagina: {section['page_width']} x {section['page_height']} cm</li>
        <li>Margini: alto {section['top']} cm, basso {section['bottom']} cm, sinistra {section['left']} cm, destra {section['right']} cm</li>
        <li>Distanze: header {section['header']} cm, footer {section['footer']} cm</li>
        <li>Media: <code>{escape(', '.join(section['media']) or '-')}</code></li>
      </ul>
    </div>
    <div class="panel">
      <h3>Da verificare col cliente</h3>
      <ul class="review-list">
        <li>Confermare posizione e stile di ogni blocco colorato.</li>
        <li>Confermare margini e uso della pagina intestata.</li>
        <li>Confermare quali placeholder generano <code>needs_review</code>.</li>
        <li>Confermare se le regole sono valide anche per altre lettere dello studio.</li>
      </ul>
    </div>
  </section>

  <h2>Header e footer ereditati dal template</h2>
  <section class="top-grid">
    <div class="panel"><h3>Header</h3><ul>{header_lines}</ul></div>
    <div class="panel"><h3>Footer</h3><ul>{footer_lines}</ul></div>
    <div class="panel"><h3>Regola</h3><p>Header/footer/logo non sono blocchi del corpo: vengono ereditati dal template e non devono essere ricreati dall'AI.</p></div>
  </section>

  <h2>Paragrafi del corpo, sezionati e misurati</h2>
  <div class="layout">
    <section class="letter-stack" aria-label="Paragrafi sezionati">
{cards}
    </section>
    <aside class="panel legend" aria-label="Legenda colori">
      <h2>Legenda colori</h2>
      {render_legend()}
      <h2>Variabili cliente</h2>
      <ul class="review-list">
        <li>Font e dimensione per ogni ruolo.</li>
        <li>Allineamento e rientri.</li>
        <li>Spazio prima/dopo e interlinea.</li>
        <li>Grassetto/corsivo per label, numeri, firma e placeholder.</li>
        <li>Regole keep-with-next / keep-together per evitare titoli isolati.</li>
        <li>Regole di review/fallback per casi non standard.</li>
      </ul>
    </aside>
  </div>

  <p style="margin-top:28px">Generato da <code>tool/scripts/build_alfa_letter_dissection_html.py</code> il {generated}.</p>
</main>
</body>
</html>
"""


def render_report() -> str:
    doc = Document(str(DOCX))
    section = section_measure(doc)
    generated = datetime.now(timezone.utc).isoformat(timespec="seconds")
    roles = Counter(ROLE_BY_INDEX[i].kind for i, _ in enumerate(doc.paragraphs))
    role_lines = "\n".join(f"- {kind}: {count}" for kind, count in sorted(roles.items()))
    return f"""# Alfa Letter Dissection Report

- generated_at: {generated}
- script: `tool/scripts/build_alfa_letter_dissection_html.py`
- input_docx: `{DOCX.relative_to(REPO_ROOT)}`
- output_html: `{OUT.relative_to(REPO_ROOT)}`
- paragraphs_measured: {len(doc.paragraphs)}
- page_cm: {section['page_width']} x {section['page_height']}
- margins_cm: top {section['top']}, bottom {section['bottom']}, left {section['left']}, right {section['right']}

## Roles

{role_lines}

## Purpose

Client-facing formatting verification page. The HTML shows every paragraph of the
Alfa synthetic letter with role, color, detection rule, Word styling values,
spacing, indentation, run-level formatting, and client confirmation notes.
"""


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(render_html(), encoding="utf-8")
    REPORT.write_text(render_report(), encoding="utf-8")
    print(f"wrote {OUT.relative_to(REPO_ROOT)}")
    print(f"wrote {REPORT.relative_to(REPO_ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
