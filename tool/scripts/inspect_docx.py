#!/usr/bin/env python3
"""Inspect a DOCX reference file for format-definition work.

Read-only. Prints page setup, headers/footers, media, and paragraph-level style
signals needed to infer deterministic formatting rules.
"""

from __future__ import annotations

import sys
from docx import Document
from docx.shared import Emu


ALIGN = {None: "-", 0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY"}


def cm(value):
    return None if value is None else round(Emu(value).cm, 2)


def pt(value):
    return None if value is None else round(value.pt, 1)


def section_info(doc):
    out = []
    for i, section in enumerate(doc.sections):
        out.append({
            "section": i,
            "page_cm": (cm(section.page_width), cm(section.page_height)),
            "margins_cm": {
                "top": cm(section.top_margin),
                "bottom": cm(section.bottom_margin),
                "left": cm(section.left_margin),
                "right": cm(section.right_margin),
                "header": cm(section.header_distance),
                "footer": cm(section.footer_distance),
            },
            "header_linked": section.header.is_linked_to_previous,
            "footer_linked": section.footer.is_linked_to_previous,
            "header_text": [p.text for p in section.header.paragraphs if p.text.strip()],
            "footer_text": [p.text for p in section.footer.paragraphs if p.text.strip()],
        })
    return out


def media_info(doc):
    return [
        str(part.partname)
        for part in doc.part.package.iter_parts()
        if "media" in str(part.partname)
    ]


def run_summary(paragraph):
    runs = []
    for run in paragraph.runs:
        font = run.font
        runs.append({
            "t": run.text,
            "font": font.name,
            "sz": pt(font.size),
            "b": font.bold,
            "i": font.italic,
            "u": font.underline,
        })
    return runs


def para_info(paragraph):
    fmt = paragraph.paragraph_format
    return {
        "style": paragraph.style.name if paragraph.style else None,
        "align": ALIGN.get(
            int(paragraph.alignment) if paragraph.alignment is not None else None,
            str(paragraph.alignment),
        ),
        "left_indent_cm": cm(fmt.left_indent),
        "right_indent_cm": cm(fmt.right_indent),
        "first_line_cm": cm(fmt.first_line_indent),
        "space_before_pt": pt(fmt.space_before),
        "space_after_pt": pt(fmt.space_after),
        "line_spacing": fmt.line_spacing,
        "keep_with_next": fmt.keep_with_next,
        "keep_together": fmt.keep_together,
        "text": paragraph.text,
        "runs": run_summary(paragraph),
    }


def inspect(path: str) -> None:
    doc = Document(path)
    print(f"### FILE: {path}")
    print("## SECTIONS")
    for section in section_info(doc):
        print(section)
    print("## MEDIA:", media_info(doc))
    print(f"## DEFAULT STYLE: {doc.styles['Normal'].font.name} {pt(doc.styles['Normal'].font.size)}")
    print(f"## PARAGRAPHS ({len(doc.paragraphs)})")
    for i, paragraph in enumerate(doc.paragraphs):
        info = para_info(paragraph)
        fonts = {(r["font"], r["sz"], r["b"], r["i"], r["u"]) for r in info["runs"]}
        runs = "" if len(fonts) <= 1 else f" RUNS={info['runs']}"
        first_run = info["runs"][0] if info["runs"] else {}
        print(
            f"[{i:02d}] {info['align']:7} li={info['left_indent_cm']} "
            f"ri={info['right_indent_cm']} fi={info['first_line_cm']} "
            f"sb={info['space_before_pt']} sa={info['space_after_pt']} "
            f"ls={info['line_spacing']} kwn={info['keep_with_next']} "
            f"kt={info['keep_together']} | {first_run} | "
            f"{info['text'][:90]!r}{runs}"
        )


def main(argv: list[str]) -> int:
    if not argv:
        raise SystemExit("usage: inspect_docx.py file.docx [file2.docx ...]")
    for path in argv:
        inspect(path)
        print()
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
