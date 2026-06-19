#!/usr/bin/env python3
"""Inspect a .docx: page setup, headers/footers, and per-paragraph formatting.

Read-only. Used to reverse-engineer the letter structure from reference outputs
and the empty template. No assumptions from external rule files.
"""
from __future__ import annotations

import sys
from docx import Document
from docx.shared import Emu

ALIGN = {None: "-", 0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY"}


def cm(v):
    return None if v is None else round(Emu(v).cm, 2)


def pt(v):
    return None if v is None else round(v.pt, 1)


def section_info(doc):
    out = []
    for i, s in enumerate(doc.sections):
        out.append({
            "section": i,
            "page_cm": (cm(s.page_width), cm(s.page_height)),
            "margins_cm": {
                "top": cm(s.top_margin), "bottom": cm(s.bottom_margin),
                "left": cm(s.left_margin), "right": cm(s.right_margin),
                "header": cm(s.header_distance), "footer": cm(s.footer_distance),
            },
            "header_linked": s.header.is_linked_to_previous,
            "footer_linked": s.footer.is_linked_to_previous,
            "header_text": [p.text for p in s.header.paragraphs if p.text.strip()],
            "footer_text": [p.text for p in s.footer.paragraphs if p.text.strip()],
        })
    return out


def media_info(doc):
    parts = doc.part.package.iter_parts()
    imgs = [p.partname for p in parts if "media" in str(p.partname)]
    return [str(x) for x in imgs]


def run_summary(p):
    runs = []
    for r in p.runs:
        f = r.font
        runs.append({
            "t": r.text,
            "font": f.name,
            "sz": pt(f.size),
            "b": f.bold,
            "i": f.italic,
        })
    return runs


def para_info(p):
    pf = p.paragraph_format
    return {
        "style": p.style.name if p.style else None,
        "align": ALIGN.get(int(p.alignment) if p.alignment is not None else None, str(p.alignment)),
        "left_indent_cm": cm(pf.left_indent),
        "right_indent_cm": cm(pf.right_indent),
        "first_line_cm": cm(pf.first_line_indent),
        "space_before_pt": pt(pf.space_before),
        "space_after_pt": pt(pf.space_after),
        "line_spacing": pf.line_spacing,
        "keep_with_next": pf.keep_with_next,
        "keep_together": pf.keep_together,
        "text": p.text,
        "runs": run_summary(p),
    }


def main(path):
    doc = Document(path)
    print(f"### FILE: {path}")
    print("## SECTIONS")
    for s in section_info(doc):
        print(s)
    print("## MEDIA:", media_info(doc))
    print(f"## DEFAULT STYLE: {doc.styles['Normal'].font.name} {pt(doc.styles['Normal'].font.size)}")
    print(f"## PARAGRAPHS ({len(doc.paragraphs)})")
    for i, p in enumerate(doc.paragraphs):
        info = para_info(p)
        # compact one-line per paragraph; runs only if mixed formatting
        fonts = {(r["font"], r["sz"], r["b"], r["i"]) for r in info["runs"]}
        rtag = "" if len(fonts) <= 1 else f" RUNS={info['runs']}"
        print(f"[{i:02d}] {info['align']:7} li={info['left_indent_cm']} "
              f"sb={info['space_before_pt']} sa={info['space_after_pt']} "
              f"ls={info['line_spacing']} kwn={info['keep_with_next']} "
              f"| {info['runs'][0] if info['runs'] else {}} "
              f"| {info['text'][:90]!r}{rtag}")


if __name__ == "__main__":
    for arg in sys.argv[1:]:
        main(arg)
        print()
