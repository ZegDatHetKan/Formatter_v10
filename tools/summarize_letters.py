#!/usr/bin/env python3
"""Compact structural summary of letter outputs to reverse-engineer rules.

For each paragraph prints: alignment, dominant font size, left-indent, spacing,
bold/italic of first run, keep_with_next, and a coarse text shape tag.
Read-only.
"""
from __future__ import annotations

import re
import sys
from collections import Counter
from docx import Document
from docx.shared import Emu

ALIGN = {None: "-", 0: "L", 1: "C", 2: "R", 3: "J"}


def cm(v):
    return None if v is None else round(Emu(v).cm, 2)


def pt(v):
    return None if v is None else round(v.pt, 1)


def dominant_size(p):
    c = Counter(pt(r.font.size) for r in p.runs if r.text.strip())
    return c.most_common(1)[0][0] if c else None


def shape(text):
    t = text.strip()
    if not t:
        return "EMPTY"
    tags = []
    if re.match(r"^\d+[.)]\s", t):
        tags.append("NUM")
    if re.match(r"^(OGGETTO|Oggetto)\s*:", t):
        tags.append("OGG")
    letters = [c for c in t if c.isalpha()]
    if letters and all(c.isupper() for c in letters) and len(t) < 60:
        tags.append("UPPER")
    if "[DA INSERIRE" in t or "[" in t and "]" in t:
        tags.append("PLH")
    if len(t) < 45:
        tags.append("short")
    return ",".join(tags) or "text"


def main(path):
    doc = Document(path)
    print(f"\n===== {path.split('/')[-1]} =====")
    sizes = Counter()
    aligns = Counter()
    indents = Counter()
    for p in doc.paragraphs:
        sizes[dominant_size(p)] += 1
        aligns[ALIGN.get(int(p.alignment) if p.alignment is not None else None)] += 1
        indents[cm(p.paragraph_format.left_indent)] += 1
    print(f"sizes={dict(sizes)} aligns={dict(aligns)} left_indents={dict(indents)}")
    for i, p in enumerate(doc.paragraphs):
        pf = p.paragraph_format
        r0 = p.runs[0].font if p.runs else None
        b = "B" if (r0 and r0.bold) else "."
        it = "I" if (r0 and r0.italic) else "."
        print(f"[{i:02d}] {ALIGN.get(int(p.alignment) if p.alignment is not None else None)} "
              f"sz={dominant_size(p)} li={cm(pf.left_indent)} "
              f"sb={pt(pf.space_before)} sa={pt(pf.space_after)} "
              f"{b}{it} kwn={'1' if pf.keep_with_next else '.'} "
              f"{shape(p.text):16} {p.text.strip()[:60]!r}")


if __name__ == "__main__":
    for a in sys.argv[1:]:
        main(a)
